from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_detailed_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('實證迴歸模型詳細統計分析報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. 假說 1 與 假說 2 分析
    doc.add_heading('一、 核心假說與交互作用分析', level=1)
    doc.add_paragraph('根據「實證模型結果_修正後」與「實證模型2結果_修正後」：')
    doc.add_paragraph('AI 技術領先：AI 變數係數約為 0.0797 (p=0.002)，在所有模型中均保持穩定的極顯著正向影響。', style='List Bullet')
    doc.add_paragraph('TESG 績效：TESG 分數係數約為 0.0688 (p=0.018)，顯示良好的永續發展表現對報酬率有顯著提振。', style='List Bullet')
    doc.add_paragraph('集團調節效應：在假說 2 中，「槓桿比與集團」的交互作用項係數為 -0.0707 (p=0.015)，顯示集團背景會顯著削弱財務槓桿對報酬的正面影響。', style='List Bullet')

    # 2. 維度分組模型分析
    doc.add_heading('二、 維度 1、2、3 之加法模型分析', level=1)
    doc.add_paragraph('分析維度對報酬率的直接影響：')
    doc.add_paragraph('維度 1 (事件衝擊)：呈現負向影響 (-0.0315, p=0.096)，顯示事件衝擊對企業股價有負面壓力。', style='List Bullet')
    doc.add_paragraph('維度 2 & 3：在這兩個分組中，財務與治理變數的顯著性與主模型高度一致，AI 與 TESG 依然是最強的預測因子。', style='List Bullet')

    # 3. 假說 2 下的維度交互作用
    doc.add_heading('三、 假說 2 架構下的維度交互作用分析', level=1)
    doc.add_paragraph('探討集團背景如何影響不同維度的反應：')
    doc.add_paragraph('一致性發現：在所有維度的交互作用模型中，「槓桿比:集團」始終保持顯著負向 (p < 0.05)，驗證了集團企業在資金運用邏輯上與獨立企業的顯著差異。', style='List Bullet')
    doc.add_paragraph('穩健性：即便加入不同維度控制項，AI 變數的係數仍維持在 0.082 左右且極顯著。', style='List Bullet')

    # 4. 統計結論
    doc.add_heading('四、 統計結論與研究建議', level=1)
    doc.add_paragraph('本研究實證結果強烈支持「AI 技術」與「ESG 績效」是提升企業異常報酬率的關鍵動能。此外，研究發現集團控制型態在財務槓桿的運用上具有顯著的調節作用，這對於投資者在評估具有集團背景企業的財務策略時具有重要參考價值。')

    # Save
    doc.save('Empirical_Analysis_Detailed_Report.docx')
    print("Successfully created Empirical_Analysis_Detailed_Report.docx")

if __name__ == "__main__":
    create_detailed_report()
